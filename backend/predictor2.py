import os
from pydantic import BaseModel
from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
import pandas as pd
from sentence_transformers import CrossEncoder, util
from huggingface_hub import InferenceClient
import numpy as np
import re
import psycopg2
from dotenv import load_dotenv
import requests
from docx import Document

# -------------------------------
# PERFORMANCE / RESOURCE CONTROL
# -------------------------------
# Limit CPU threads to avoid overconsumption (important for servers / containers)
os.environ["OPENBLAS_NUM_THREADS"] = "3"
os.environ["OMP_NUM_THREADS"] = "3"
os.environ["MKL_NUM_THREADS"] = "3"

# -------------------------------
# HUGGINGFACE CONFIGURATION
# -------------------------------
# Disable telemetry and UI noise (clean logs, privacy, performance)os.environ["HF_TOKEN"] = "xNxeIgnPSPRuBreOnKTMgQOAGMTS"
os.environ["HF_HUB_DISABLE_TELEMETRY"] = "1"
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["HF_HUB_DISABLE_PROGRESS_BARS"] = "1"

class CandidateBot:
    def __init__(self):
        load_dotenv()

        self.hf_token = os.getenv("HF_TOKEN")
        if not self.hf_token:
            raise ValueError("HF_TOKEN not loaded from environment.")

        self.client = InferenceClient(token=self.hf_token)
        self.llm_client = InferenceClient(token=self.hf_token, provider="featherless-ai")
        self.embed_model = 'sentence-transformers/all-mpnet-base-v2'
        self.rerank_model = 'cross-encoder/ms-marco-MiniLM-L-6-v2'
        self.reranker = CrossEncoder(self.rerank_model)
        self.llm_model = 'meta-llama/Llama-3.1-8B-Instruct'
        
        self.headers = {"Authorization": f"Bearer {self.hf_token}"}

        # Load data 
        self.dataframe = self.getDataFrame()

        # Prepare embeddings 
        self.prepare_search_data( )

    def call_model_api(self, question):
        response = self.client.feature_extraction(
            model=self.embed_model,
            text=question,
            normalize=True
        )
        return np.array(response)

        #embeddings = np.array(embeddings)
        #norms = np.linalg.norm(embeddings, axis=1, keepdims=True)
        #return embeddings / np.clip(norms, 1e-10, None)

    def call_reranker_api(self, query, candidates):
        pairs = [[query, candidate] for candidate in candidates]
        return np.array(self.reranker.predict(pairs))

    def call_llm(self, prompt):
        try:
            response = self.llm_client.chat_completion(
                messages=[
                    {"role": "user", "content": prompt},
                ],
                model=self.llm_model,
                max_tokens=300,
                temperature=0.13,
            )
        except Exception as e:
            raise ValueError(f"Error calling LLM: {e}")
        return response.choices[0].message.content.strip()
    
    def call_llm_val_question(self, prompt):
        try:
            response = self.llm_client.chat_completion(
                messages=[
                    {"role": "user", "content": prompt},
                ],
                model=self.llm_model,
                max_tokens=150,
                temperature=0.13,
            )
        except Exception as e:
            raise ValueError(f"Error calling LLM: {e}")
        return response.choices[0].message.content.strip()

    # -------------------------------
    # TEXT CLEANING
    # -------------------------------
    def clean_text(self, text):
        text = str(text).lower()            # Convert to string + lowercase
        text = re.sub(r"\s+", " ", text)    # Replace multiple spaces with single space
        return text.strip()                 # Remove leading/trailing spaces


    def get_connection(self):
        #Get DB connection using psycopg2 and .env configuration
        load_dotenv()
        db_url = os.getenv("DB_URL")
        if not db_url:
            raise ValueError("DB_URL not loaded from .env")
        return psycopg2.connect(db_url)

    # -------------------------------
    # LOAD DATAFRAME
    # -------------------------------
    def getDataFrame(self):
        #Get data from PostgreSQL database and return as DataFrame
        #Could be from SAP HANA Database
        conn = self.get_connection()
        try:
            with conn.cursor() as cur:
                cur.execute( f""" SELECT question, answer FROM myprofile; """ )
                rows = cur.fetchall()
                columns = ["question", "answer"]
                return pd.DataFrame(rows, columns=columns)
            conn.close()
        except Exception as e:
            return f"Error reading data: {e}"

    # -------------------------------
    # PREPARE SEARCH DATA
    # -------------------------------
    def prepare_search_data(self):

        # Create combined searchable text
        self.dataframe["search_text"] = (
            "question: " + self.dataframe["question"].astype(str) +
            " answer: " + self.dataframe["answer"].astype(str)
        )
        self.questions = self.dataframe["search_text"].tolist()

        if self.questions:
            # Prepare embeddings
            # Convert all Q&A into vector representations (semantic space)
            self.q_emb = self.call_model_api(self.questions)
        else:
            raise ValueError("No questions found in the dataframe to prepare search data.")
        
    # -------------------------------
    # MAIN SEARCH FUNCTION
    # -------------------------------
    def process_question(self, user_input, top_k=5):

        # Clean user input
        user_input = self.clean_text(user_input)
        if not user_input:
            return "Please provide a valid question."

        # Convert user question into embedding
        u_emb = self.call_model_api([user_input])

        # Compute cosine similarity between user question and all stored questions
        sims = util.dot_score(u_emb, self.q_emb)[0].cpu().numpy()

        # Get top-k most similar results
        top_k = min(top_k, len(sims))
        top_idx = np.argsort(sims)[::-1][:top_k]

        # Best similarity score
        best_retrieval_score = float(sims[top_idx[0]])
        #print(f"1 - score: {best_retrieval_score:.4f}")

        # First confidence filter
        RETRIEVAL_THRESHOLD = 0.25
        if best_retrieval_score < RETRIEVAL_THRESHOLD:
            return "Sorry, I couldn't find a reliable answer. Could you please rephrase or provide more details?"
        
        # Rerank the top results for better accuracy
        pairs_texts = [self.questions[i] for i in top_idx]
        rerank_scores = self.call_reranker_api(user_input, pairs_texts)

        # Select best reranked result
        best_local_idx = int(np.argmax(rerank_scores))
        best_idx = int(top_idx[best_local_idx])
        best_rerank_score = float(rerank_scores[best_local_idx])
        #print(f"2 - score: {best_rerank_score:.4f}")
        
        #best_rerank_score = -1

        RERANK_THRESHOLD = 0.00
        if best_rerank_score < RERANK_THRESHOLD:
            new_answer = self.get_new_answer(user_input)
            if new_answer and new_answer != "question not relevant":
                self.save_new_qa(user_input, new_answer, autogen=True)
                return new_answer
            else:
                return "Sorry, I couldn't find a reliable answer. Could you please rephrase or provide more details?"
        
        return self.dataframe.iloc[best_idx]["answer"]
    
    def get_new_answer(self, question):
        # Read the profile.docx file 
        profile_doc = Document('data/profile.docx')
        profile_content = '\n'.join([para.text for para in profile_doc.paragraphs if para.text.strip()])

        # Read the sapecosystem.docx file 
        sapecosystem_doc = Document('data/sapecosystem.docx')
        sapecosystem_content = '\n'.join([parag.text for parag in sapecosystem_doc.paragraphs if parag.text.strip()])

        prompt = (
            f"You are a SAP Expert hiring assistant. Execute the following steps:\n"
            f"1. Based on the SAPECOSYSTEM content and PROFILE content, validate if the QUESTION is inside the SAPECOSYSTEM and PROFILE context.\n"
            f"2. If QUESTION not inside the context of SAPECOSYSTEM, return 'question not relevant' else return blank. Don't mention anything else, PROFILE or QUESTION.\n"
            f"SAPECOSYSTEM:\n{sapecosystem_content}\n\n"
            f"PROFILE:\n{profile_content}\n\n"
            f"QUESTION:\n{question}\n\n"
        )
        question_validation = self.call_llm_val_question(prompt)
        #print(f"3 - question_validation: {question_validation}")

        if question_validation.strip().lower() != "question not relevant":
            prompt = (
                f"You are a SAP Expert hiring assistant. Execute the following steps:\n"
                f"1. Based on the PROFILE content, analyse the QUESTION, validate in candidate's PROFILE if he knows the subject asked in the QUESTION.\n"
                f"2. If candidate doesn't know, respond he needs to improve this part and he is open to learn new things.\n"
                f"3. Refers to the candidate as 'he'.\n"
                f"Give a precise answer to the question, in 150 characters maximum. Don't mention PROFILE or QUESTION.\n"
                f"PROFILE:\n{profile_content}\n\n"
                f"QUESTION:\n{question}\n\n"
            )
            return self.call_llm(prompt)
        else:
            return "question not relevant"

    def save_new_qa(self, question, answer, autogen=False, ip=""):
        # Save new Q&A pair to PostgreSQL database
        conn = self.get_connection()
        try:
            with conn.cursor() as cur:
                insert_sql = f""" INSERT INTO myprofile (question, answer, autogen, ip) VALUES (%s, %s, %s, %s); """
                cur.execute(insert_sql, (question, answer, autogen, ip))
            conn.commit()
        finally:
            conn.close()
